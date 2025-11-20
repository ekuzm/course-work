from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING

NBSP = "\u00A0"  # неразрывный пробел


def set_font(run, name: str, bold: bool = False):
    run.font.name = name
    run.font.size = Pt(14)
    run.font.bold = bold


def make_header_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = None
    pf.first_line_indent = None
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(0)
    r = p.add_run(text)
    set_font(r, "Times New Roman")
    return p


def make_normal_paragraph(doc: Document, text: str):
    """Обычный абзац с красной строкой 1.25 см, TNR 14."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = None
    pf.first_line_indent = Cm(1.25)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(0)
    r = p.add_run(text)
    set_font(r, "Times New Roman")
    return p


def make_class_heading(doc: Document, number: int, kind: str, name: str):
    """
    Заголовок: "N Класс Name:" или "N Структура Name:"
    N + 'Класс/Структура' — Times New Roman
    Name — Courier New
    : — Times New Roman
    """
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = None
    pf.first_line_indent = Cm(1.25)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(0)

    r1 = p.add_run(f"{number} {kind} ")
    set_font(r1, "Times New Roman")

    r2 = p.add_run(name)
    set_font(r2, "Courier New")

    r3 = p.add_run(":")
    set_font(r3, "Times New Roman")

    return p


def add_bullet_block(doc: Document, title: str, items: list[dict]):
    """
    items: список словарей вида {"code": "...", "desc": "..."}.
    code — шрифт кода (Courier New), desc — Times New Roman.
    """
    # Заголовок "Поля:" / "Методы:"
    make_normal_paragraph(doc, title)

    if not items:
        return

    n = len(items)
    for idx, item in enumerate(items):
        last = (idx == n - 1)
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Cm(2.5)
        pf.first_line_indent = None
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.space_after = Pt(0)

        code = item.get("code", "").strip()
        desc = (item.get("desc", "") or "").strip()

        # Маркер "– " (с NBSP), Times New Roman
        r0 = p.add_run("–" + NBSP)
        set_font(r0, "Times New Roman")

        # Код — Courier New
        if code:
            r_code = p.add_run(code)
            set_font(r_code, "Courier New")

        # Описание — Times New Roman
        ending = "." if last else ";"

        if desc:
            r_colon = p.add_run(": ")
            set_font(r_colon, "Times New Roman")

            r_desc = p.add_run(desc + ending)
            set_font(r_desc, "Times New Roman")
        else:
            # Нет описания — просто код + точка/точка с запятой
            r_end = p.add_run(ending)
            set_font(r_end, "Times New Roman")


def build_sections() -> list[dict]:
    """
    Возвращает список всех сущностей (классы/структуры) проекта управления IT-компанией.
    """
    sections: list[dict] = []

    # 1. Employee (базовый класс)
    sections.append({
        "kind": "Класс",
        "name": "Employee",
        "desc": "Базовый абстрактный класс сотрудника компании с общими характеристиками и методами.",
        "fields": [
            {"code": "int id", "desc": "уникальный идентификатор сотрудника"},
            {"code": "QString name", "desc": "имя сотрудника"},
            {"code": "QString position", "desc": "должность сотрудника"},
            {"code": "double salary", "desc": "месячная зарплата сотрудника"},
            {"code": "QString department", "desc": "отдел сотрудника"},
            {"code": "bool isActive", "desc": "флаг активности сотрудника (не уволен)"},
            {"code": "double employmentRate", "desc": "коэффициент занятости (0.0-1.0)"},
            {"code": "int weeklyHoursCapacity", "desc": "недельная емкость в часах"},
            {"code": "int currentWeeklyHours", "desc": "текущая недельная загрузка в часах"},
            {"code": "std::vector<int> assignedProjects", "desc": "список ID назначенных проектов"},
            {"code": "std::vector<int> projectHistory", "desc": "история участия в проектах"},
        ],
        "methods": [
            {"code": "Employee(int employeeId, QString name, QString position, double salary, QString department, double employmentRate = 1.0, int weeklyCapacity = 40)", "desc": "конструктор сотрудника"},
            {"code": "virtual ~Employee() = default", "desc": "виртуальный деструктор"},
            {"code": "virtual QString getEmployeeType() const = 0", "desc": "получить тип сотрудника (абстрактный метод)"},
            {"code": "virtual QString getDetails() const", "desc": "получить детальную информацию о сотруднике"},
            {"code": "virtual double calculateBonus() const = 0", "desc": "рассчитать бонус сотрудника (абстрактный метод)"},
            {"code": "int getId() const", "desc": "получить идентификатор"},
            {"code": "QString getName() const", "desc": "получить имя"},
            {"code": "QString getPosition() const", "desc": "получить должность"},
            {"code": "double getSalary() const", "desc": "получить зарплату"},
            {"code": "QString getDepartment() const", "desc": "получить отдел"},
            {"code": "bool getIsActive() const", "desc": "проверить активность"},
            {"code": "double getEmploymentRate() const", "desc": "получить коэффициент занятости"},
            {"code": "int getWeeklyHoursCapacity() const", "desc": "получить недельную емкость"},
            {"code": "int getCurrentWeeklyHours() const", "desc": "получить текущую загрузку"},
            {"code": "bool isAvailable(int requestedHours) const", "desc": "проверить доступность для назначения часов"},
            {"code": "int getAvailableHours() const", "desc": "получить доступные часы"},
            {"code": "void addWeeklyHours(int hours)", "desc": "добавить часы к загрузке"},
            {"code": "void removeWeeklyHours(int hours)", "desc": "удалить часы из загрузки"},
            {"code": "const std::vector<int>& getAssignedProjects() const", "desc": "получить список назначенных проектов"},
            {"code": "void addAssignedProject(int projectId)", "desc": "добавить проект в список назначений"},
            {"code": "void removeAssignedProject(int projectId)", "desc": "удалить проект из списка назначений"},
            {"code": "bool isAssignedToProject(int projectId) const", "desc": "проверить назначение на проект"},
            {"code": "const std::vector<int>& getProjectHistory() const", "desc": "получить историю проектов"},
            {"code": "void addToProjectHistory(int projectId)", "desc": "добавить проект в историю"},
            {"code": "void setIsActive(bool active)", "desc": "установить статус активности"},
            {"code": "void setEmploymentRate(double newRate, int baseWeeklyCapacity = 40)", "desc": "установить коэффициент занятости"},
        ],
    })

    # 2. Manager
    sections.append({
        "kind": "Класс",
        "name": "Manager",
        "desc": "Класс менеджера, наследующий Employee, с дополнительным полем управляемого проекта.",
        "fields": [
            {"code": "int managedProjectId", "desc": "идентификатор управляемого проекта"},
        ],
        "methods": [
            {"code": "Manager(int employeeId, QString employeeName, double employeeSalary, QString employeeDepartment, int managedProjectId = -1, double employmentRate = 1.0)", "desc": "конструктор менеджера"},
            {"code": "QString getEmployeeType() const override", "desc": "возвращает \"Manager\""},
            {"code": "QString getDetails() const override", "desc": "получить детальную информацию о менеджере"},
            {"code": "double calculateBonus() const override", "desc": "рассчитать бонус менеджера"},
            {"code": "int getManagedProjectId() const", "desc": "получить ID управляемого проекта"},
            {"code": "void setManagedProjectId(int projectId)", "desc": "установить управляемый проект"},
        ],
    })

    # 3. Developer
    sections.append({
        "kind": "Класс",
        "name": "Developer",
        "desc": "Класс разработчика с информацией о языке программирования и опыте работы.",
        "fields": [
            {"code": "QString programmingLanguage", "desc": "язык программирования"},
            {"code": "double yearsOfExperience", "desc": "годы опыта работы"},
        ],
        "methods": [
            {"code": "Developer(int employeeId, QString employeeName, double employeeSalary, QString employeeDepartment, QString developerProgrammingLanguage, double developerYearsOfExperience, double employmentRate = 1.0)", "desc": "конструктор разработчика"},
            {"code": "QString getEmployeeType() const override", "desc": "возвращает \"Developer\""},
            {"code": "QString getDetails() const override", "desc": "получить детальную информацию о разработчике"},
            {"code": "double calculateBonus() const override", "desc": "рассчитать бонус разработчика"},
            {"code": "QString getProgrammingLanguage() const", "desc": "получить язык программирования"},
            {"code": "double getYearsOfExperience() const", "desc": "получить годы опыта"},
        ],
    })

    # 4. Designer
    sections.append({
        "kind": "Класс",
        "name": "Designer",
        "desc": "Класс дизайнера с информацией об инструменте дизайна и количестве проектов.",
        "fields": [
            {"code": "QString designTool", "desc": "инструмент дизайна"},
            {"code": "int numberOfProjects", "desc": "количество выполненных проектов"},
        ],
        "methods": [
            {"code": "Designer(int employeeId, QString employeeName, double employeeSalary, QString employeeDepartment, QString designerTool, int designerNumberOfProjects, double employmentRate = 1.0)", "desc": "конструктор дизайнера"},
            {"code": "QString getEmployeeType() const override", "desc": "возвращает \"Designer\""},
            {"code": "QString getDetails() const override", "desc": "получить детальную информацию о дизайнере"},
            {"code": "double calculateBonus() const override", "desc": "рассчитать бонус дизайнера"},
            {"code": "QString getDesignTool() const", "desc": "получить инструмент дизайна"},
            {"code": "int getNumberOfProjects() const", "desc": "получить количество проектов"},
        ],
    })

    # 5. QA
    sections.append({
        "kind": "Класс",
        "name": "QA",
        "desc": "Класс тестировщика с информацией о типе тестирования и количестве найденных багов.",
        "fields": [
            {"code": "QString testingType", "desc": "тип тестирования"},
            {"code": "int bugsFound", "desc": "количество найденных багов"},
        ],
        "methods": [
            {"code": "QA(int employeeId, QString employeeName, double employeeSalary, QString employeeDepartment, QString qaTestingType, int qaBugsFound, double employmentRate = 1.0)", "desc": "конструктор тестировщика"},
            {"code": "QString getEmployeeType() const override", "desc": "возвращает \"QA\""},
            {"code": "QString getDetails() const override", "desc": "получить детальную информацию о тестировщике"},
            {"code": "double calculateBonus() const override", "desc": "рассчитать бонус тестировщика"},
            {"code": "QString getTestingType() const", "desc": "получить тип тестирования"},
            {"code": "int getBugsFound() const", "desc": "получить количество найденных багов"},
        ],
    })

    # 6. Task
    sections.append({
        "kind": "Класс",
        "name": "Task",
        "desc": "Класс задачи проекта с информацией о типе, приоритете, оцененных и выделенных часах.",
        "fields": [
            {"code": "int id", "desc": "идентификатор задачи"},
            {"code": "QString name", "desc": "название задачи"},
            {"code": "QString type", "desc": "тип задачи (Management, Development, Design, QA)"},
            {"code": "int estimatedHours", "desc": "оцененные часы выполнения"},
            {"code": "int allocatedHours", "desc": "выделенные часы"},
            {"code": "int priority", "desc": "приоритет задачи"},
            {"code": "QString phase", "desc": "фаза выполнения (Planned, In Progress, Completed)"},
        ],
        "methods": [
            {"code": "Task(int taskId, const QString& name, const QString& type, int estimatedHours, int priority)", "desc": "конструктор задачи"},
            {"code": "int getId() const", "desc": "получить идентификатор"},
            {"code": "QString getName() const", "desc": "получить название"},
            {"code": "QString getType() const", "desc": "получить тип"},
            {"code": "int getEstimatedHours() const", "desc": "получить оцененные часы"},
            {"code": "int getAllocatedHours() const", "desc": "получить выделенные часы"},
            {"code": "int getPriority() const", "desc": "получить приоритет"},
            {"code": "QString getPhase() const", "desc": "получить фазу"},
            {"code": "void setEstimatedHours(int hours)", "desc": "установить оцененные часы"},
            {"code": "void setAllocatedHours(int hours)", "desc": "установить выделенные часы"},
            {"code": "void addAllocatedHours(int hours)", "desc": "добавить выделенные часы"},
            {"code": "void setPhase(const QString& phaseValue)", "desc": "установить фазу"},
        ],
    })

    # 7. ProjectParams
    sections.append({
        "kind": "Структура",
        "name": "ProjectParams",
        "desc": "Параметры для создания проекта.",
        "fields": [
            {"code": "int projectId", "desc": "идентификатор проекта"},
            {"code": "const QString& name", "desc": "название проекта"},
            {"code": "const QString& description", "desc": "описание проекта"},
            {"code": "const QString& phase", "desc": "фаза жизненного цикла"},
            {"code": "const QDate& startDate", "desc": "дата начала"},
            {"code": "const QDate& endDate", "desc": "дата окончания"},
            {"code": "double budget", "desc": "бюджет проекта"},
            {"code": "const QString& clientName", "desc": "имя клиента"},
            {"code": "int estimatedHours", "desc": "оцененные часы"},
        ],
        "methods": [
            {"code": "(нет)", "desc": "структура только для передачи параметров"},
        ],
    })

    # 8. Project
    sections.append({
        "kind": "Класс",
        "name": "Project",
        "desc": "Класс проекта компании с задачами, бюджетом и информацией о жизненном цикле.",
        "fields": [
            {"code": "int id", "desc": "идентификатор проекта"},
            {"code": "QString name", "desc": "название проекта"},
            {"code": "QString description", "desc": "описание проекта"},
            {"code": "QString phase", "desc": "фаза жизненного цикла (Analysis, Planning, Design, Development, Testing, Deployment, Maintenance, Completed)"},
            {"code": "QDate startDate", "desc": "дата начала проекта"},
            {"code": "QDate endDate", "desc": "дата окончания проекта"},
            {"code": "double budget", "desc": "бюджет проекта"},
            {"code": "QString clientName", "desc": "имя клиента"},
            {"code": "int initialEstimatedHours", "desc": "начальные оцененные часы"},
            {"code": "int allocatedHours", "desc": "выделенные часы"},
            {"code": "double employeeCosts", "desc": "стоимость сотрудников"},
            {"code": "std::vector<Task> tasks", "desc": "список задач проекта"},
        ],
        "methods": [
            {"code": "explicit Project(const ProjectParams& params)", "desc": "конструктор проекта"},
            {"code": "int getId() const", "desc": "получить идентификатор"},
            {"code": "QString getName() const", "desc": "получить название"},
            {"code": "QString getDescription() const", "desc": "получить описание"},
            {"code": "QString getPhase() const", "desc": "получить фазу"},
            {"code": "QDate getStartDate() const", "desc": "получить дату начала"},
            {"code": "QDate getEndDate() const", "desc": "получить дату окончания"},
            {"code": "double getBudget() const", "desc": "получить бюджет"},
            {"code": "QString getClientName() const", "desc": "получить имя клиента"},
            {"code": "int getEstimatedHours() const", "desc": "получить оцененные часы"},
            {"code": "int getInitialEstimatedHours() const", "desc": "получить начальные оцененные часы"},
            {"code": "int getAllocatedHours() const", "desc": "получить выделенные часы"},
            {"code": "double getEmployeeCosts() const", "desc": "получить стоимость сотрудников"},
            {"code": "const std::vector<Task>& getTasks() const", "desc": "получить список задач (const)"},
            {"code": "std::vector<Task>& getTasks()", "desc": "получить список задач"},
            {"code": "void addTask(const Task& task)", "desc": "добавить задачу в проект"},
            {"code": "void clearTasks()", "desc": "очистить список задач"},
            {"code": "int getTasksEstimatedTotal() const", "desc": "получить сумму оцененных часов всех задач"},
            {"code": "int getTasksAllocatedTotal() const", "desc": "получить сумму выделенных часов всех задач"},
            {"code": "int getNextTaskId() const", "desc": "получить следующий ID задачи"},
            {"code": "static int getPhaseOrder(const QString& phaseName)", "desc": "получить порядок фазы в жизненном цикле"},
            {"code": "void setPhase(const QString& newPhase)", "desc": "установить фазу проекта"},
            {"code": "void setBudget(double newBudget)", "desc": "установить бюджет"},
            {"code": "void setEstimatedHours(int hours)", "desc": "установить оцененные часы"},
            {"code": "void setAllocatedHours(int hours)", "desc": "установить выделенные часы"},
            {"code": "void addEmployeeCost(double cost)", "desc": "добавить стоимость сотрудника"},
            {"code": "void removeEmployeeCost(double cost)", "desc": "удалить стоимость сотрудника"},
            {"code": "void recomputeTotalsFromTasks()", "desc": "пересчитать итоги из задач"},
            {"code": "bool isActive() const", "desc": "проверить активность проекта"},
        ],
    })

    # 9. EmployeeContainer
    sections.append({
        "kind": "Класс",
        "name": "EmployeeContainer",
        "desc": "Контейнер для хранения сотрудников компании.",
        "fields": [
            {"code": "std::vector<std::shared_ptr<Employee>> employees", "desc": "вектор указателей на сотрудников"},
        ],
        "methods": [
            {"code": "void add(std::shared_ptr<Employee> employee)", "desc": "добавить сотрудника"},
            {"code": "void remove(int employeeId)", "desc": "удалить сотрудника по ID"},
            {"code": "std::shared_ptr<Employee> find(int employeeId) const", "desc": "найти сотрудника по ID"},
            {"code": "std::vector<std::shared_ptr<Employee>> getAll() const", "desc": "получить всех сотрудников"},
            {"code": "size_t size() const", "desc": "получить количество сотрудников"},
        ],
    })

    # 10. ProjectContainer
    sections.append({
        "kind": "Класс",
        "name": "ProjectContainer",
        "desc": "Контейнер для хранения проектов компании.",
        "fields": [
            {"code": "std::vector<std::shared_ptr<Project>> projects", "desc": "вектор указателей на проекты"},
        ],
        "methods": [
            {"code": "void add(std::shared_ptr<Project> project)", "desc": "добавить проект"},
            {"code": "void remove(int projectId)", "desc": "удалить проект по ID"},
            {"code": "std::shared_ptr<Project> find(int projectId) const", "desc": "найти проект по ID"},
            {"code": "std::vector<std::shared_ptr<Project>> getAll() const", "desc": "получить все проекты"},
            {"code": "size_t size() const", "desc": "получить количество проектов"},
        ],
    })

    # 11. TaskAssignmentManager
    sections.append({
        "kind": "Класс",
        "name": "TaskAssignmentManager",
        "desc": "Менеджер назначений сотрудников на задачи с управлением часами и пересчетом.",
        "fields": [
            {"code": "std::map<std::tuple<int, int, int>, int>& taskAssignments", "desc": "карта назначений (employeeId, projectId, taskId) -> hours"},
            {"code": "EmployeeContainer& employees", "desc": "контейнер сотрудников"},
            {"code": "ProjectContainer& projects", "desc": "контейнер проектов"},
        ],
        "methods": [
            {"code": "TaskAssignmentManager(std::map<std::tuple<int, int, int>, int>& assignments, EmployeeContainer& empContainer, ProjectContainer& projContainer)", "desc": "конструктор менеджера"},
            {"code": "void assignEmployeeToTask(int employeeId, int projectId, int taskId, int hours)", "desc": "назначить сотрудника на задачу"},
            {"code": "void restoreTaskAssignment(int employeeId, int projectId, int taskId, int hours)", "desc": "восстановить назначение"},
            {"code": "void removeEmployeeTaskAssignments(int employeeId)", "desc": "удалить все назначения сотрудника"},
            {"code": "void recalculateEmployeeHours() const", "desc": "пересчитать часы сотрудников"},
            {"code": "void recalculateTaskAllocatedHours() const", "desc": "пересчитать выделенные часы задач"},
            {"code": "void fixTaskAssignmentsToCapacity()", "desc": "исправить назначения до емкости сотрудников"},
            {"code": "void autoAssignEmployeesToProject(int projectId)", "desc": "автоматически назначить сотрудников на проект"},
            {"code": "int getEmployeeProjectHours(int employeeId, int projectId) const", "desc": "получить часы сотрудника на проект"},
            {"code": "int getEmployeeTaskHours(int employeeId, int projectId, int taskId) const", "desc": "получить часы сотрудника на задачу"},
            {"code": "void scaleEmployeeTaskAssignments(int employeeId, double scaleFactor)", "desc": "масштабировать назначения сотрудника"},
            {"code": "int getTaskAssignment(int employeeId, int projectId, int taskId) const", "desc": "получить назначение"},
            {"code": "void setTaskAssignment(int employeeId, int projectId, int taskId, int hours)", "desc": "установить назначение"},
            {"code": "void addTaskAssignment(int employeeId, int projectId, int taskId, int hours)", "desc": "добавить назначение"},
            {"code": "void removeTaskAssignment(int employeeId, int projectId, int taskId)", "desc": "удалить назначение"},
            {"code": "std::map<std::tuple<int, int, int>, int> getAllTaskAssignments() const", "desc": "получить все назначения"},
        ],
    })

    # 12. CompanyStatistics
    sections.append({
        "kind": "Класс",
        "name": "CompanyStatistics",
        "desc": "Класс для расчета статистики компании по сотрудникам и проектам.",
        "fields": [
            {"code": "const EmployeeContainer& employees", "desc": "контейнер сотрудников"},
            {"code": "const ProjectContainer& projects", "desc": "контейнер проектов"},
        ],
        "methods": [
            {"code": "CompanyStatistics(const EmployeeContainer& empContainer, const ProjectContainer& projContainer)", "desc": "конструктор статистики"},
            {"code": "int getEmployeeCount() const", "desc": "получить количество сотрудников"},
            {"code": "int getProjectCount() const", "desc": "получить количество проектов"},
            {"code": "double getTotalSalaries() const", "desc": "получить сумму всех зарплат"},
            {"code": "double getTotalBudget() const", "desc": "получить сумму всех бюджетов проектов"},
        ],
    })

    # 13. Company
    sections.append({
        "kind": "Класс",
        "name": "Company",
        "desc": "Основной класс компании, содержащий сотрудников, проекты и управление назначениями.",
        "fields": [
            {"code": "QString name", "desc": "название компании"},
            {"code": "QString industry", "desc": "отрасль компании"},
            {"code": "QString location", "desc": "местоположение компании"},
            {"code": "int foundedYear", "desc": "год основания"},
            {"code": "EmployeeContainer employees", "desc": "контейнер сотрудников"},
            {"code": "mutable ProjectContainer projects", "desc": "контейнер проектов"},
            {"code": "std::map<std::tuple<int, int, int>, int> taskAssignments", "desc": "карта назначений"},
            {"code": "TaskAssignmentManager taskManager", "desc": "менеджер назначений"},
            {"code": "mutable CompanyStatistics statistics", "desc": "статистика компании"},
        ],
        "methods": [
            {"code": "Company(QString name, QString industry, QString location, int foundedYear)", "desc": "конструктор компании"},
            {"code": "Company(Company&& other) noexcept", "desc": "конструктор перемещения"},
            {"code": "~Company() noexcept", "desc": "деструктор"},
            {"code": "TaskAssignmentManager& getTaskManager()", "desc": "получить менеджер назначений"},
            {"code": "const CompanyStatistics& getStatistics() const", "desc": "получить статистику"},
            {"code": "QString getName() const", "desc": "получить название"},
            {"code": "QString getIndustry() const", "desc": "получить отрасль"},
            {"code": "QString getLocation() const", "desc": "получить местоположение"},
            {"code": "int getFoundedYear() const", "desc": "получить год основания"},
            {"code": "void addEmployee(std::shared_ptr<Employee> employee)", "desc": "добавить сотрудника"},
            {"code": "void removeEmployee(int employeeId)", "desc": "удалить сотрудника"},
            {"code": "std::shared_ptr<Employee> getEmployee(int employeeId) const", "desc": "получить сотрудника по ID"},
            {"code": "std::vector<std::shared_ptr<Employee>> getAllEmployees() const", "desc": "получить всех сотрудников"},
            {"code": "void addProject(const Project& project)", "desc": "добавить проект"},
            {"code": "void removeProject(int projectId)", "desc": "удалить проект"},
            {"code": "Project* getProject(int projectId) const", "desc": "получить проект по ID"},
            {"code": "std::vector<Project> getAllProjects() const", "desc": "получить все проекты"},
            {"code": "void addTaskToProject(int projectId, const Task& task) const", "desc": "добавить задачу в проект"},
            {"code": "std::vector<Task> getProjectTasks(int projectId) const", "desc": "получить задачи проекта"},
            {"code": "void assignEmployeeToTask(int employeeId, int projectId, int taskId, int hours)", "desc": "назначить сотрудника на задачу"},
            {"code": "void restoreTaskAssignment(int employeeId, int projectId, int taskId, int hours)", "desc": "восстановить назначение"},
            {"code": "void removeEmployeeTaskAssignments(int employeeId)", "desc": "удалить назначения сотрудника"},
            {"code": "void recalculateEmployeeHours()", "desc": "пересчитать часы сотрудников"},
            {"code": "void recalculateTaskAllocatedHours()", "desc": "пересчитать выделенные часы задач"},
            {"code": "void fixTaskAssignmentsToCapacity()", "desc": "исправить назначения до емкости"},
            {"code": "void recalculateAllHours()", "desc": "пересчитать все часы"},
            {"code": "void autoAssignEmployeesToProject(int projectId)", "desc": "автоматически назначить сотрудников на проект"},
            {"code": "int getEmployeeHours(int employeeId, int projectId, int taskId = -1) const", "desc": "получить часы сотрудника"},
            {"code": "void scaleEmployeeTaskAssignments(int employeeId, double scaleFactor)", "desc": "масштабировать назначения сотрудника"},
            {"code": "std::map<std::tuple<int, int, int>, int> getAllTaskAssignments() const", "desc": "получить все назначения"},
            {"code": "int getEmployeeCount() const", "desc": "получить количество сотрудников"},
            {"code": "int getProjectCount() const", "desc": "получить количество проектов"},
            {"code": "double getTotalSalaries() const", "desc": "получить сумму зарплат"},
            {"code": "double getTotalBudget() const", "desc": "получить сумму бюджетов"},
            {"code": "QString getCompanyInfo() const", "desc": "получить информацию о компании"},
        ],
    })

    # 14. CostCalculationService
    sections.append({
        "kind": "Класс",
        "name": "CostCalculationService",
        "desc": "Сервис расчета стоимости сотрудников на основе зарплаты и часов.",
        "fields": [
            {"code": "(нет)", "desc": "содержит только статические методы"},
        ],
        "methods": [
            {"code": "static double calculateHourlyRate(double monthlySalary)", "desc": "рассчитать часовую ставку из месячной зарплаты"},
            {"code": "static double calculateEmployeeCost(double monthlySalary, int hours)", "desc": "рассчитать стоимость сотрудника за часы"},
        ],
    })

    # 15. TaskAssignmentService
    sections.append({
        "kind": "Класс",
        "name": "TaskAssignmentService",
        "desc": "Сервис назначения сотрудников на задачи с валидацией и автоматическим распределением.",
        "fields": [
            {"code": "Company* company", "desc": "указатель на компанию"},
        ],
        "methods": [
            {"code": "explicit TaskAssignmentService(Company* company)", "desc": "конструктор сервиса"},
            {"code": "void assignEmployeeToTask(int employeeId, int projectId, int taskId, int hours)", "desc": "назначить сотрудника на задачу"},
            {"code": "void restoreTaskAssignment(int employeeId, int projectId, int taskId, int hours)", "desc": "восстановить назначение"},
            {"code": "void removeEmployeeTaskAssignments(int employeeId)", "desc": "удалить назначения сотрудника"},
            {"code": "void scaleEmployeeTaskAssignments(int employeeId, double scaleFactor)", "desc": "масштабировать назначения"},
            {"code": "void fixTaskAssignmentsToCapacity()", "desc": "исправить назначения до емкости"},
            {"code": "void recalculateTaskAllocatedHours()", "desc": "пересчитать выделенные часы"},
            {"code": "void autoAssignEmployeesToProject(int projectId)", "desc": "автоматически назначить сотрудников на проект"},
            {"code": "int getEmployeeProjectHours(int employeeId, int projectId) const", "desc": "получить часы сотрудника на проект"},
            {"code": "int getEmployeeTaskHours(int employeeId, int projectId, int taskId) const", "desc": "получить часы сотрудника на задачу"},
            {"code": "static bool roleMatchesSDLCStage(const QString& employeePosition, const QString& projectPhase)", "desc": "проверить соответствие роли фазе SDLC"},
            {"code": "static bool taskTypeMatchesEmployeeType(const QString& taskType, const QString& employeeType)", "desc": "проверить соответствие типа задачи типу сотрудника"},
        ],
    })

    # 16. EmployeeService
    sections.append({
        "kind": "Класс",
        "name": "EmployeeService",
        "desc": "Сервис для работы с сотрудниками и расчета их загрузки.",
        "fields": [
            {"code": "mutable Company* company", "desc": "указатель на компанию"},
        ],
        "methods": [
            {"code": "explicit EmployeeService(Company* company)", "desc": "конструктор сервиса"},
            {"code": "void recalculateEmployeeHours() const", "desc": "пересчитать часы сотрудников"},
            {"code": "int calculateTotalAssignedHours(int employeeId) const", "desc": "рассчитать общие назначенные часы сотрудника"},
        ],
    })

    # 17. ProjectService
    sections.append({
        "kind": "Класс",
        "name": "ProjectService",
        "desc": "Сервис для работы с проектами и задачами.",
        "fields": [
            {"code": "(нет)", "desc": "содержит только статические методы"},
        ],
        "methods": [
            {"code": "static void addTaskToProject(const Company* company, int projectId, const Task& task)", "desc": "добавить задачу в проект"},
            {"code": "static std::vector<Task> getProjectTasks(const Company* company, int projectId)", "desc": "получить задачи проекта"},
            {"code": "static void recomputeProjectTotals(const Company* company, int projectId)", "desc": "пересчитать итоги проекта"},
        ],
    })

    # 18. BaseException
    sections.append({
        "kind": "Класс",
        "name": "BaseException",
        "desc": "Базовый класс исключений проекта, наследующий std::exception.",
        "fields": [
            {"code": "QString message", "desc": "сообщение об ошибке"},
        ],
        "methods": [
            {"code": "explicit BaseException(QString msg)", "desc": "конструктор исключения"},
            {"code": "const char* what() const noexcept override", "desc": "получить сообщение об ошибке"},
            {"code": "QString getMessage() const", "desc": "получить сообщение как QString"},
        ],
    })

    # 19. EmployeeException
    sections.append({
        "kind": "Класс",
        "name": "EmployeeException",
        "desc": "Исключение для ошибок работы с сотрудниками.",
        "fields": [
            {"code": "(нет)", "desc": "наследует BaseException"},
        ],
        "methods": [
            {"code": "explicit EmployeeException(QString msg)", "desc": "конструктор исключения"},
        ],
    })

    # 20. CompanyException
    sections.append({
        "kind": "Класс",
        "name": "CompanyException",
        "desc": "Исключение для ошибок работы с компанией.",
        "fields": [
            {"code": "(нет)", "desc": "наследует BaseException"},
        ],
        "methods": [
            {"code": "explicit CompanyException(QString msg)", "desc": "конструктор исключения"},
        ],
    })

    # 21. ProjectException
    sections.append({
        "kind": "Класс",
        "name": "ProjectException",
        "desc": "Исключение для ошибок работы с проектами.",
        "fields": [
            {"code": "(нет)", "desc": "наследует BaseException"},
        ],
        "methods": [
            {"code": "explicit ProjectException(QString msg)", "desc": "конструктор исключения"},
        ],
    })

    # 22. TaskException
    sections.append({
        "kind": "Класс",
        "name": "TaskException",
        "desc": "Исключение для ошибок работы с задачами.",
        "fields": [
            {"code": "(нет)", "desc": "наследует BaseException"},
        ],
        "methods": [
            {"code": "explicit TaskException(QString msg)", "desc": "конструктор исключения"},
        ],
    })

    # 23. FileManagerException
    sections.append({
        "kind": "Класс",
        "name": "FileManagerException",
        "desc": "Исключение для ошибок работы с файлами.",
        "fields": [
            {"code": "(нет)", "desc": "наследует BaseException"},
        ],
        "methods": [
            {"code": "explicit FileManagerException(QString msg)", "desc": "конструктор исключения"},
        ],
    })

    # 24. ExceptionHandler
    sections.append({
        "kind": "Класс",
        "name": "ExceptionHandler",
        "desc": "Обработчик исключений для отображения ошибок пользователю.",
        "fields": [
            {"code": "(нет)", "desc": "содержит только статические методы"},
        ],
        "methods": [
            {"code": "static void handleCompanyException(const CompanyException& e, QDialog* dialog, const QString& action)", "desc": "обработать исключение компании"},
            {"code": "static void handleFileManagerException(const FileManagerException& e, QDialog* dialog, const QString& action)", "desc": "обработать исключение файлового менеджера"},
            {"code": "static void handleGenericException(const std::exception& e, QDialog* dialog)", "desc": "обработать общее исключение"},
        ],
    })

    # 25. FileManager
    sections.append({
        "kind": "Класс",
        "name": "FileManager",
        "desc": "Менеджер для сохранения и загрузки данных компании в текстовом формате.",
        "fields": [
            {"code": "static std::map<int, bool> employeeStatusesFromFile", "desc": "статусы сотрудников из файла"},
        ],
        "methods": [
            {"code": "static void saveCompany(const Company& company, const QString& fileName)", "desc": "сохранить компанию"},
            {"code": "static void saveEmployees(const Company& company, const QString& fileName)", "desc": "сохранить сотрудников"},
            {"code": "static void saveProjects(const Company& company, const QString& fileName)", "desc": "сохранить проекты"},
            {"code": "static void saveTasks(const Company& company, const QString& fileName)", "desc": "сохранить задачи"},
            {"code": "static void saveTaskAssignments(const Company& company, const QString& fileName)", "desc": "сохранить назначения"},
            {"code": "static Company loadCompany(const QString& fileName)", "desc": "загрузить компанию"},
            {"code": "static void loadEmployees(Company& company, const QString& fileName)", "desc": "загрузить сотрудников"},
            {"code": "static void loadProjects(Company& company, const QString& fileName)", "desc": "загрузить проекты"},
            {"code": "static void loadTasks(Company& company, const QString& fileName)", "desc": "загрузить задачи"},
            {"code": "static void loadTaskAssignments(Company& company, const QString& fileName)", "desc": "загрузить назначения"},
            {"code": "static Company loadFromFile(const QString& fileName)", "desc": "загрузить компанию из файла"},
        ],
    })

    # 26. AutoSaveLoader
    sections.append({
        "kind": "Класс",
        "name": "AutoSaveLoader",
        "desc": "Менеджер автоматического сохранения и загрузки данных приложения.",
        "fields": [
            {"code": "(нет)", "desc": "содержит только статические методы"},
        ],
        "methods": [
            {"code": "static QString getDataDirectory()", "desc": "получить директорию данных"},
            {"code": "static void autoSave(const std::vector<Company*>& companies, MainWindow* mainWindow)", "desc": "автоматически сохранить данные"},
            {"code": "static void autoLoad(std::vector<Company*>& companies, Company*& currentCompany, int& currentCompanyIndex, MainWindow* mainWindow)", "desc": "автоматически загрузить данные"},
            {"code": "static void clearDataFiles(const QString& dataDirPath)", "desc": "очистить файлы данных"},
        ],
    })

    # 27. CompanyManager
    sections.append({
        "kind": "Класс",
        "name": "CompanyManager",
        "desc": "Менеджер для управления несколькими компаниями в приложении.",
        "fields": [
            {"code": "(нет)", "desc": "содержит только статические методы"},
        ],
        "methods": [
            {"code": "static void initializeCompany(std::span<Company* const> companies, Company*& currentCompany, const int& currentCompanyIndex, int& nextEmployeeId, int& nextProjectId, QComboBox* selector)", "desc": "инициализировать компанию"},
            {"code": "static void addCompany(std::vector<Company*>& companies, Company*& currentCompany, int& currentCompanyIndex, QComboBox* selector, QWidget* parent)", "desc": "добавить компанию"},
            {"code": "static void switchCompany(std::vector<Company*>& companies, Company*& currentCompany, int& currentCompanyIndex, QComboBox* selector, int newIndex)", "desc": "переключить компанию"},
            {"code": "static void deleteCompany(std::vector<Company*>& companies, Company*& currentCompany, int& currentCompanyIndex, QComboBox* selector, QWidget* parent)", "desc": "удалить компанию"},
            {"code": "static void refreshCompanyList(std::span<Company* const> companies, QComboBox* selector)", "desc": "обновить список компаний"},
        ],
    })

    # 28. EmployeeTabUI
    sections.append({
        "kind": "Структура",
        "name": "EmployeeTabUI",
        "desc": "Структура для хранения UI-компонентов вкладки сотрудников.",
        "fields": [
            {"code": "QWidget* tab", "desc": "виджет вкладки"},
            {"code": "QTableWidget* table", "desc": "таблица сотрудников"},
            {"code": "QPushButton* addBtn", "desc": "кнопка добавления"},
            {"code": "QLineEdit* searchEdit", "desc": "поле поиска"},
        ],
        "methods": [
            {"code": "(нет)", "desc": "структура только для хранения указателей"},
        ],
    })

    # 29. ProjectTabUI
    sections.append({
        "kind": "Структура",
        "name": "ProjectTabUI",
        "desc": "Структура для хранения UI-компонентов вкладки проектов.",
        "fields": [
            {"code": "QWidget* tab", "desc": "виджет вкладки"},
            {"code": "QWidget* listContainer", "desc": "контейнер списка проектов"},
            {"code": "QWidget* detailHeaderContainer", "desc": "контейнер заголовка деталей"},
            {"code": "QWidget* detailContainer", "desc": "контейнер деталей проекта"},
            {"code": "QTableWidget* table", "desc": "таблица проектов"},
            {"code": "QTableWidget* tasksTable", "desc": "таблица задач"},
            {"code": "QPushButton* addBtn", "desc": "кнопка добавления проекта"},
            {"code": "QPushButton* detailCloseBtn", "desc": "кнопка закрытия деталей"},
            {"code": "QPushButton* detailAutoAssignBtn", "desc": "кнопка автоматического назначения"},
            {"code": "QLabel* detailTitle", "desc": "заголовок деталей"},
            {"code": "QTextEdit* detailInfoText", "desc": "текст информации о проекте"},
        ],
        "methods": [
            {"code": "(нет)", "desc": "структура только для хранения указателей"},
        ],
    })

    # 30. StatisticsTabUI
    sections.append({
        "kind": "Структура",
        "name": "StatisticsTabUI",
        "desc": "Структура для хранения UI-компонентов вкладки статистики.",
        "fields": [
            {"code": "QWidget* tab", "desc": "виджет вкладки"},
            {"code": "QWidget* chartWidget", "desc": "виджет диаграммы"},
            {"code": "QWidget* chartInnerWidget", "desc": "внутренний виджет диаграммы"},
            {"code": "QTextEdit* text", "desc": "текстовое поле статистики"},
        ],
        "methods": [
            {"code": "(нет)", "desc": "структура только для хранения указателей"},
        ],
    })

    # 31. CompanyUI
    sections.append({
        "kind": "Структура",
        "name": "CompanyUI",
        "desc": "Структура для хранения UI-компонентов управления компаниями.",
        "fields": [
            {"code": "QWidget* widget", "desc": "виджет управления"},
            {"code": "QComboBox* selector", "desc": "селектор компаний"},
            {"code": "QPushButton* addBtn", "desc": "кнопка добавления компании"},
            {"code": "QPushButton* deleteBtn", "desc": "кнопка удаления компании"},
        ],
        "methods": [
            {"code": "(нет)", "desc": "структура только для хранения указателей"},
        ],
    })

    # 32. MainWindow
    sections.append({
        "kind": "Класс",
        "name": "MainWindow",
        "desc": "Главное окно приложения с вкладками для управления сотрудниками, проектами и статистикой.",
        "fields": [
            {"code": "QTabWidget* mainTabWidget", "desc": "виджет вкладок"},
            {"code": "EmployeeTabUI employeeUI", "desc": "UI вкладки сотрудников"},
            {"code": "ProjectTabUI projectUI", "desc": "UI вкладки проектов"},
            {"code": "StatisticsTabUI statisticsUI", "desc": "UI вкладки статистики"},
            {"code": "CompanyUI companyUI", "desc": "UI управления компаниями"},
            {"code": "std::vector<Company*> companies", "desc": "вектор компаний"},
            {"code": "Company* currentCompany", "desc": "текущая компания"},
            {"code": "int currentCompanyIndex", "desc": "индекс текущей компании"},
            {"code": "int nextEmployeeId", "desc": "следующий ID сотрудника"},
            {"code": "int nextProjectId", "desc": "следующий ID проекта"},
            {"code": "int detailedProjectId", "desc": "ID детализированного проекта"},
            {"code": "int pendingTaskSelectionId", "desc": "ID ожидающей выбора задачи"},
        ],
        "methods": [
            {"code": "explicit MainWindow(QWidget* parent = nullptr)", "desc": "конструктор главного окна"},
            {"code": "~MainWindow() override", "desc": "деструктор"},
            {"code": "void refreshAllData()", "desc": "обновить все данные"},
            {"code": "void autoSave()", "desc": "автоматически сохранить данные"},
            {"code": "void selectProjectRowById(int projectId)", "desc": "выбрать строку проекта по ID"},
            {"code": "void closeEvent(QCloseEvent* event) override", "desc": "обработчик закрытия окна"},
            {"code": "static QString getDataDirectory()", "desc": "получить директорию данных"},
            {"code": "void validateAndFixProjectAssignments(const Company* company)", "desc": "валидировать и исправить назначения"},
        ],
    })

    # 33. StatisticsChartWidget
    sections.append({
        "kind": "Класс",
        "name": "StatisticsChartWidget",
        "desc": "Виджет для отображения статистики компании с анимированными диаграммами.",
        "fields": [
            {"code": "const Company* company", "desc": "указатель на компанию"},
            {"code": "QTimer* animationTimer", "desc": "таймер анимации"},
            {"code": "double animationProgress", "desc": "прогресс анимации"},
        ],
        "methods": [
            {"code": "explicit StatisticsChartWidget(QWidget* parent = nullptr)", "desc": "конструктор виджета"},
            {"code": "void setData(const Company* companyData)", "desc": "установить данные компании"},
            {"code": "void paintEvent(QPaintEvent* event) override", "desc": "обработчик отрисовки"},
        ],
    })

    # 34. SafeValue
    sections.append({
        "kind": "Класс",
        "name": "SafeValue",
        "desc": "Шаблонный класс для безопасного хранения значений с проверкой диапазона.",
        "fields": [
            {"code": "T value", "desc": "значение"},
            {"code": "T minValue", "desc": "минимальное значение"},
            {"code": "T maxValue", "desc": "максимальное значение"},
            {"code": "bool isValid", "desc": "флаг валидности значения"},
        ],
        "methods": [
            {"code": "SafeValue(T val, T min, T max)", "desc": "конструктор с проверкой диапазона"},
            {"code": "bool isValidValue() const", "desc": "проверить валидность"},
            {"code": "T getValue() const", "desc": "получить значение (или минимум если невалидно)"},
            {"code": "T getClampedValue() const", "desc": "получить значение, ограниченное диапазоном"},
            {"code": "void setValue(T newValue)", "desc": "установить значение"},
        ],
    })

    # 35. CompanyTaskOperations
    sections.append({
        "kind": "Класс",
        "name": "CompanyTaskOperations",
        "desc": "Класс операций с назначениями задач для компании (обертка над TaskAssignmentManager).",
        "fields": [
            {"code": "(нет)", "desc": "содержит только статические методы"},
        ],
        "methods": [
            {"code": "static void assignEmployeeToTask(Company* company, int employeeId, int projectId, int taskId, int hours)", "desc": "назначить сотрудника на задачу"},
            {"code": "static void restoreTaskAssignment(Company* company, int employeeId, int projectId, int taskId, int hours)", "desc": "восстановить назначение"},
            {"code": "static void removeEmployeeTaskAssignments(Company* company, int employeeId)", "desc": "удалить назначения сотрудника"},
            {"code": "static void recalculateEmployeeHours(Company* company)", "desc": "пересчитать часы сотрудников"},
            {"code": "static void recalculateTaskAllocatedHours(Company* company)", "desc": "пересчитать выделенные часы"},
            {"code": "static void fixTaskAssignmentsToCapacity(Company* company)", "desc": "исправить назначения до емкости"},
            {"code": "static void autoAssignEmployeesToProject(Company* company, int projectId)", "desc": "автоматически назначить сотрудников на проект"},
            {"code": "static int getEmployeeProjectHours(const Company* company, int employeeId, int projectId)", "desc": "получить часы сотрудника на проект"},
            {"code": "static int getEmployeeTaskHours(const Company* company, int employeeId, int projectId, int taskId)", "desc": "получить часы сотрудника на задачу"},
            {"code": "static void scaleEmployeeTaskAssignments(Company* company, int employeeId, double scaleFactor)", "desc": "масштабировать назначения"},
            {"code": "static int getTaskAssignment(const Company* company, int employeeId, int projectId, int taskId)", "desc": "получить назначение"},
            {"code": "static void setTaskAssignment(Company* company, int employeeId, int projectId, int taskId, int hours)", "desc": "установить назначение"},
            {"code": "static void addTaskAssignment(Company* company, int employeeId, int projectId, int taskId, int hours)", "desc": "добавить назначение"},
            {"code": "static void removeTaskAssignment(Company* company, int employeeId, int projectId, int taskId)", "desc": "удалить назначение"},
            {"code": "static std::map<std::tuple<int, int, int>, int> getAllTaskAssignments(const Company* company)", "desc": "получить все назначения"},
        ],
    })

    # 36. EmployeeOperations
    sections.append({
        "kind": "Класс",
        "name": "EmployeeOperations",
        "desc": "Класс операций с сотрудниками в главном окне.",
        "fields": [
            {"code": "(нет)", "desc": "содержит только статические методы"},
        ],
        "methods": [
            {"code": "static void addEmployee(MainWindow* window)", "desc": "добавить сотрудника"},
            {"code": "static void editEmployee(MainWindow* window)", "desc": "редактировать сотрудника"},
            {"code": "static void deleteEmployee(MainWindow* window)", "desc": "удалить сотрудника"},
            {"code": "static void fireEmployee(MainWindow* window)", "desc": "уволить сотрудника"},
            {"code": "static void searchEmployee(MainWindow* window)", "desc": "поиск сотрудника"},
            {"code": "static void refreshEmployeeTable(MainWindow* window)", "desc": "обновить таблицу сотрудников"},
        ],
    })

    # 37. ProjectOperations
    sections.append({
        "kind": "Класс",
        "name": "ProjectOperations",
        "desc": "Класс операций с проектами в главном окне.",
        "fields": [
            {"code": "(нет)", "desc": "содержит только статические методы"},
        ],
        "methods": [
            {"code": "static void addProject(MainWindow* window)", "desc": "добавить проект"},
            {"code": "static void editProject(MainWindow* window)", "desc": "редактировать проект"},
            {"code": "static void deleteProject(MainWindow* window)", "desc": "удалить проект"},
            {"code": "static void refreshProjectTable(MainWindow* window)", "desc": "обновить таблицу проектов"},
            {"code": "static void openProjectDetails(MainWindow* window)", "desc": "открыть детали проекта"},
            {"code": "static void closeProjectDetails(MainWindow* window)", "desc": "закрыть детали проекта"},
            {"code": "static void autoAssignDetailedProject(MainWindow* window)", "desc": "автоматически назначить на детализированный проект"},
            {"code": "static void assignTaskFromDetails(MainWindow* window, int projectId = -1, int taskId = -1)", "desc": "назначить задачу из деталей"},
            {"code": "static void addProjectTask(MainWindow* window)", "desc": "добавить задачу в проект"},
            {"code": "static void assignEmployeeToTask(MainWindow* window)", "desc": "назначить сотрудника на задачу"},
            {"code": "static void autoAssignToProject(MainWindow* window, int projectId = -1)", "desc": "автоматически назначить на проект"},
            {"code": "static void viewProjectAssignments(MainWindow* window)", "desc": "просмотреть назначения проекта"},
            {"code": "static void viewEmployeeHistory(MainWindow* window)", "desc": "просмотреть историю сотрудника"},
            {"code": "static void showStatistics(MainWindow* window)", "desc": "показать статистику"},
        ],
    })

    # 38. ProjectDetailOperations
    sections.append({
        "kind": "Класс",
        "name": "ProjectDetailOperations",
        "desc": "Класс операций с деталями проекта.",
        "fields": [
            {"code": "(нет)", "desc": "содержит только статические методы"},
        ],
        "methods": [
            {"code": "static void showProjectDetails(MainWindow* window, int projectId)", "desc": "показать детали проекта"},
            {"code": "static void hideProjectDetails(MainWindow* window)", "desc": "скрыть детали проекта"},
            {"code": "static void refreshProjectDetailView(MainWindow* window)", "desc": "обновить вид деталей проекта"},
            {"code": "static void populateProjectTasksTable(MainWindow* window, const Project& project)", "desc": "заполнить таблицу задач проекта"},
        ],
    })

    # 39. CompanyOperations
    sections.append({
        "kind": "Класс",
        "name": "CompanyOperations",
        "desc": "Класс операций с компаниями в главном окне.",
        "fields": [
            {"code": "(нет)", "desc": "содержит только статические методы"},
        ],
        "methods": [
            {"code": "static void addCompany(MainWindow* window)", "desc": "добавить компанию"},
            {"code": "static void switchCompany(MainWindow* window)", "desc": "переключить компанию"},
            {"code": "static void deleteCompany(MainWindow* window)", "desc": "удалить компанию"},
            {"code": "static void refreshCompanyList(MainWindow* window)", "desc": "обновить список компаний"},
            {"code": "static void initializeCompanySetup(MainWindow* window)", "desc": "инициализировать настройку компании"},
        ],
    })

    # 40. MainWindowSelectionHelper
    sections.append({
        "kind": "Класс",
        "name": "MainWindowSelectionHelper",
        "desc": "Вспомогательный класс для получения выбранных элементов в главном окне.",
        "fields": [
            {"code": "(нет)", "desc": "содержит только статические методы"},
        ],
        "methods": [
            {"code": "static int getSelectedEmployeeId(const MainWindow* window)", "desc": "получить ID выбранного сотрудника"},
            {"code": "static int getSelectedProjectId(const MainWindow* window)", "desc": "получить ID выбранного проекта"},
        ],
    })

    # 41. ProjectEditData
    sections.append({
        "kind": "Структура",
        "name": "ProjectEditData",
        "desc": "Структура для хранения данных редактирования проекта.",
        "fields": [
            {"code": "QString projectName", "desc": "название проекта"},
            {"code": "double projectBudget", "desc": "бюджет проекта"},
            {"code": "QString selectedPhase", "desc": "выбранная фаза"},
            {"code": "QString clientName", "desc": "имя клиента"},
            {"code": "int estimatedHours", "desc": "оцененные часы"},
            {"code": "QString newName", "desc": "новое название"},
            {"code": "QString newDescription", "desc": "новое описание"},
            {"code": "QString newPhase", "desc": "новая фаза"},
            {"code": "QDate newStartDate", "desc": "новая дата начала"},
            {"code": "QDate newEndDate", "desc": "новая дата окончания"},
            {"code": "QString newClientName", "desc": "новое имя клиента"},
        ],
        "methods": [
            {"code": "(нет)", "desc": "структура только для хранения данных"},
        ],
    })

    # 42. MainWindowProjectDialogHandler
    sections.append({
        "kind": "Класс",
        "name": "MainWindowProjectDialogHandler",
        "desc": "Обработчик диалогов проектов в главном окне.",
        "fields": [
            {"code": "(нет)", "desc": "содержит только статические методы"},
        ],
        "methods": [
            {"code": "static void handleEditProjectDialog(MainWindow* window, int projectId, QDialog& dialog, const ProjectDialogHelper::ProjectDialogFields& fields)", "desc": "обработать диалог редактирования проекта"},
            {"code": "static void handleAddTaskDialog(MainWindow* window, int projectId, QDialog& dialog, const QLineEdit* taskNameEdit, const QComboBox* taskTypeCombo, const QLineEdit* taskEstHoursEdit, const QLineEdit* priorityEdit)", "desc": "обработать диалог добавления задачи"},
            {"code": "static void handleAssignEmployeeToTaskDialog(MainWindow* window, int projectId, QDialog& dialog, const QComboBox* taskCombo, const QComboBox* employeeCombo, const QLineEdit* hoursEdit, const std::vector<Task>& tasks)", "desc": "обработать диалог назначения сотрудника на задачу"},
        ],
    })

    # 43. MainWindowProjectDetailHelper
    sections.append({
        "kind": "Класс",
        "name": "MainWindowProjectDetailHelper",
        "desc": "Вспомогательный класс для работы с деталями проекта.",
        "fields": [
            {"code": "(нет)", "desc": "содержит только статические методы"},
        ],
        "methods": [
            {"code": "static void showProjectDetails(MainWindow* window, int projectId)", "desc": "показать детали проекта"},
            {"code": "static void hideProjectDetails(MainWindow* window)", "desc": "скрыть детали проекта"},
            {"code": "static void refreshProjectDetailView(MainWindow* window)", "desc": "обновить вид деталей"},
            {"code": "static void populateProjectTasksTable(MainWindow* window, const Project& project)", "desc": "заполнить таблицу задач"},
        ],
    })

    # 44. MainWindowTaskAssignmentHelper
    sections.append({
        "kind": "Класс",
        "name": "MainWindowTaskAssignmentHelper",
        "desc": "Вспомогательный класс для работы с назначениями задач.",
        "fields": [
            {"code": "(нет)", "desc": "содержит только статические методы"},
        ],
        "methods": [
            {"code": "static void removeEmployeeFromProjectTasks(const MainWindow* window, int employeeId, int projectId, Project* mutableProject, double employeeHourlyRate)", "desc": "удалить сотрудника из задач проекта"},
            {"code": "static void collectTaskAssignments(const MainWindow* window, int projectId, const std::vector<Task>& savedTasks, std::vector<std::tuple<int, int, int, int>>& savedTaskAssignments)", "desc": "собрать назначения задач"},
            {"code": "static void handleEmployeeActiveAssignments(const MainWindow* window, int employeeId, const std::shared_ptr<Employee>& employee)", "desc": "обработать активные назначения сотрудника"},
        ],
    })

    # 45. MainWindowValidationHelper
    sections.append({
        "kind": "Класс",
        "name": "MainWindowValidationHelper",
        "desc": "Вспомогательный класс для валидации данных в главном окне.",
        "fields": [
            {"code": "(нет)", "desc": "содержит только статические методы"},
        ],
        "methods": [
            {"code": "static bool checkCompanyAndHandleError(MainWindow* window, const QString& actionName)", "desc": "проверить наличие компании и обработать ошибку"},
            {"code": "static bool checkDuplicateProjectOnEdit(const QString& projectName, int excludeId, const Company* currentCompany)", "desc": "проверить дубликат проекта при редактировании"},
            {"code": "static void validateAndFixProjectAssignments(MainWindow* window, const Company* company)", "desc": "валидировать и исправить назначения проекта"},
        ],
    })

    # 46. MainWindowDataOperations
    sections.append({
        "kind": "Класс",
        "name": "MainWindowDataOperations",
        "desc": "Класс операций с данными в главном окне.",
        "fields": [
            {"code": "(нет)", "desc": "содержит только статические методы"},
        ],
        "methods": [
            {"code": "static void refreshAllData(MainWindow* window)", "desc": "обновить все данные"},
            {"code": "static void autoSave(MainWindow* window)", "desc": "автоматически сохранить данные"},
            {"code": "static void autoLoad(MainWindow* window)", "desc": "автоматически загрузить данные"},
            {"code": "static void selectProjectRowById(MainWindow* window, int projectId)", "desc": "выбрать строку проекта по ID"},
        ],
    })

    # 47. MainWindowUIHelper
    sections.append({
        "kind": "Класс",
        "name": "MainWindowUIHelper",
        "desc": "Вспомогательный класс для настройки UI главного окна.",
        "fields": [
            {"code": "(нет)", "desc": "содержит только статические методы"},
        ],
        "methods": [
            {"code": "static void setupUI(MainWindow* window)", "desc": "настроить UI главного окна"},
            {"code": "static void setupEmployeeTab(MainWindow* window)", "desc": "настроить вкладку сотрудников"},
            {"code": "static void setupProjectTab(MainWindow* window)", "desc": "настроить вкладку проектов"},
            {"code": "static void setupStatisticsTab(MainWindow* window)", "desc": "настроить вкладку статистики"},
            {"code": "static void drawStatisticsChart(MainWindow* window, QWidget* widget)", "desc": "нарисовать диаграмму статистики"},
            {"code": "static void clearAllDataFiles(MainWindow* window)", "desc": "очистить все файлы данных"},
            {"code": "static void setupTableWidget(QTableWidget* table, const QStringList& headers, const QList<int>& columnWidths, bool stretchLast = true)", "desc": "настроить виджет таблицы"},
            {"code": "static QWidget* createEmployeeActionButtons(MainWindow* window, int rowIndex)", "desc": "создать кнопки действий для сотрудника"},
            {"code": "static QWidget* createProjectActionButtons(MainWindow* window, int rowIndex)", "desc": "создать кнопки действий для проекта"},
            {"code": "static void selectProjectRowById(MainWindow* window, int projectId)", "desc": "выбрать строку проекта по ID"},
        ],
    })

    # 48. MainWindowUIBuilder
    sections.append({
        "kind": "Класс",
        "name": "MainWindowUIBuilder",
        "desc": "Класс для построения UI главного окна.",
        "fields": [
            {"code": "(нет)", "desc": "содержит только статические методы"},
        ],
        "methods": [
            {"code": "static void setupMainUI(MainWindow* window)", "desc": "настроить основной UI"},
            {"code": "static void setupEmployeeTab(MainWindow* window, QTabWidget* tabWidget)", "desc": "настроить вкладку сотрудников"},
            {"code": "static void setupProjectTab(MainWindow* window, QTabWidget* tabWidget)", "desc": "настроить вкладку проектов"},
            {"code": "static void setupStatisticsTab(MainWindow* window, QTabWidget* tabWidget)", "desc": "настроить вкладку статистики"},
            {"code": "static QWidget* createCompanyWidget(MainWindow* window)", "desc": "создать виджет компании"},
        ],
    })

    # 49. ProjectDialogFields
    sections.append({
        "kind": "Структура",
        "name": "ProjectDialogFields",
        "desc": "Структура для хранения полей диалога проекта.",
        "fields": [
            {"code": "QLineEdit* nameEdit", "desc": "поле ввода названия"},
            {"code": "QTextEdit* descEdit", "desc": "поле ввода описания"},
            {"code": "QComboBox* phaseCombo", "desc": "комбобокс фазы"},
            {"code": "QDateEdit* startDate", "desc": "редактор даты начала"},
            {"code": "QDateEdit* endDate", "desc": "редактор даты окончания"},
            {"code": "QLineEdit* budgetEdit", "desc": "поле ввода бюджета"},
            {"code": "QLineEdit* estimatedHoursEdit", "desc": "поле ввода оцененных часов"},
            {"code": "QLineEdit* clientNameEdit", "desc": "поле ввода имени клиента"},
            {"code": "QLabel* clientNameLabel", "desc": "лейбл имени клиента"},
            {"code": "QLineEdit* clientIndustryEdit", "desc": "поле ввода отрасли клиента"},
            {"code": "QLabel* clientIndustryLabel", "desc": "лейбл отрасли клиента"},
            {"code": "QLineEdit* clientContactEdit", "desc": "поле ввода контакта клиента"},
            {"code": "QLabel* clientContactLabel", "desc": "лейбл контакта клиента"},
            {"code": "QComboBox* projectTypeCombo", "desc": "комбобокс типа проекта"},
        ],
        "methods": [
            {"code": "(нет)", "desc": "структура только для хранения указателей на виджеты"},
        ],
    })

    # 50. ProjectDialogHelper
    sections.append({
        "kind": "Класс",
        "name": "ProjectDialogHelper",
        "desc": "Вспомогательный класс для работы с диалогами проектов.",
        "fields": [
            {"code": "(нет)", "desc": "содержит только статические методы"},
        ],
        "methods": [
            {"code": "static void createProjectDialogFields(QDialog& dialog, QFormLayout* form, ProjectDialogFields& fields)", "desc": "создать поля диалога проекта"},
            {"code": "static void populateProjectDialogFields(const Project* project, ProjectDialogFields& fields)", "desc": "заполнить поля диалога данными проекта"},
            {"code": "static void setupClientFieldsVisibility(ProjectDialogFields& fields)", "desc": "настроить видимость полей клиента"},
        ],
    })

    # 51. CreateEmployeeDialogFields
    sections.append({
        "kind": "Структура",
        "name": "CreateEmployeeDialogFields",
        "desc": "Структура для хранения полей диалога создания сотрудника.",
        "fields": [
            {"code": "QComboBox*& typeCombo", "desc": "комбобокс типа сотрудника"},
            {"code": "QLineEdit*& nameEdit", "desc": "поле ввода имени"},
            {"code": "QLineEdit*& salaryEdit", "desc": "поле ввода зарплаты"},
            {"code": "QLineEdit*& deptEdit", "desc": "поле ввода отдела"},
            {"code": "QComboBox*& employmentRateCombo", "desc": "комбобокс коэффициента занятости"},
            {"code": "QComboBox*& managerProject", "desc": "комбобокс проекта менеджера"},
            {"code": "QLineEdit*& devLanguage", "desc": "поле языка программирования"},
            {"code": "QLineEdit*& devExperience", "desc": "поле опыта разработчика"},
            {"code": "QLineEdit*& designerTool", "desc": "поле инструмента дизайнера"},
            {"code": "QLineEdit*& designerProjects", "desc": "поле количества проектов дизайнера"},
            {"code": "QLineEdit*& qaTestType", "desc": "поле типа тестирования QA"},
            {"code": "QLineEdit*& qaBugs", "desc": "поле количества багов QA"},
            {"code": "QLabel*& managerProjectLabel", "desc": "лейбл проекта менеджера"},
            {"code": "QLabel*& devLanguageLabel", "desc": "лейбл языка программирования"},
            {"code": "QLabel*& devExperienceLabel", "desc": "лейбл опыта разработчика"},
            {"code": "QLabel*& designerToolLabel", "desc": "лейбл инструмента дизайнера"},
            {"code": "QLabel*& designerProjectsLabel", "desc": "лейбл проектов дизайнера"},
            {"code": "QLabel*& qaTestTypeLabel", "desc": "лейбл типа тестирования QA"},
            {"code": "QLabel*& qaBugsLabel", "desc": "лейбл багов QA"},
        ],
        "methods": [
            {"code": "(нет)", "desc": "структура только для хранения ссылок на виджеты"},
        ],
    })
    # 52. CreateEditEmployeeDialogFields
    sections.append({
        "kind": "Структура",
        "name": "CreateEditEmployeeDialogFields",
        "desc": "Структура для хранения полей диалога редактирования сотрудника.",
        "fields": [
            {"code": "QLineEdit*& nameEdit", "desc": "поле ввода имени"},
            {"code": "QLineEdit*& salaryEdit", "desc": "поле ввода зарплаты"},
            {"code": "QLineEdit*& deptEdit", "desc": "поле ввода отдела"},
            {"code": "QComboBox*& employmentRateCombo", "desc": "комбобокс коэффициента занятости"},
            {"code": "QComboBox*& managerProject", "desc": "комбобокс проекта менеджера"},
            {"code": "QLineEdit*& devLanguage", "desc": "поле языка разработчика"},
            {"code": "QLineEdit*& devExperience", "desc": "поле опыта разработчика"},
            {"code": "QLineEdit*& designerTool", "desc": "поле инструмента дизайнера"},
            {"code": "QLineEdit*& designerProjects", "desc": "поле проектов дизайнера"},
            {"code": "QLineEdit*& qaTestType", "desc": "поле типа тестирования QA"},
            {"code": "QLineEdit*& qaBugs", "desc": "поле багов QA"},
            {"code": "QLabel*& managerProjectLabel", "desc": "лейбл проекта менеджера"},
            {"code": "QLabel*& devLanguageLabel", "desc": "лейбл языка программирования"},
            {"code": "QLabel*& devExperienceLabel", "desc": "лейбл опыта разработчика"},
            {"code": "QLabel*& designerToolLabel", "desc": "лейбл инструмента дизайнера"},
            {"code": "QLabel*& designerProjectsLabel", "desc": "лейбл проектов дизайнера"},
            {"code": "QLabel*& qaTestTypeLabel", "desc": "лейбл типа тестирования QA"},
            {"code": "QLabel*& qaBugsLabel", "desc": "лейбл багов QA"},
        ],
        "methods": [
            {"code": "(нет)", "desc": "структура только для хранения ссылок на виджеты"},
        ],
    })

    # 53. SetAllFieldsVisibleParams
    sections.append({
        "kind": "Структура",
        "name": "SetAllFieldsVisibleParams",
        "desc": "Структура параметров для установки видимости всех полей сотрудника.",
        "fields": [
            {"code": "QLabel* managerProjectLabel", "desc": "лейбл проекта менеджера"},
            {"code": "QComboBox* managerProject", "desc": "комбобокс проекта менеджера"},
            {"code": "QLabel* devLanguageLabel", "desc": "лейбл языка программирования"},
            {"code": "QLineEdit* devLanguage", "desc": "поле языка разработчика"},
            {"code": "QLabel* devExperienceLabel", "desc": "лейбл опыта разработчика"},
            {"code": "QLineEdit* devExperience", "desc": "поле опыта разработчика"},
            {"code": "QLabel* designerToolLabel", "desc": "лейбл инструмента дизайнера"},
            {"code": "QLineEdit* designerTool", "desc": "поле инструмента дизайнера"},
            {"code": "QLabel* designerProjectsLabel", "desc": "лейбл проектов дизайнера"},
            {"code": "QLineEdit* designerProjects", "desc": "поле проектов дизайнера"},
            {"code": "QLabel* qaTestTypeLabel", "desc": "лейбл типа тестирования QA"},
            {"code": "QLineEdit* qaTestType", "desc": "поле типа тестирования QA"},
            {"code": "QLabel* qaBugsLabel", "desc": "лейбл багов QA"},
            {"code": "QLineEdit* qaBugs", "desc": "поле багов QA"},
        ],
        "methods": [
            {"code": "(нет)", "desc": "структура только для передачи параметров"},
        ],
    })

    # 54. PopulateEmployeeFieldsParams
    sections.append({
        "kind": "Структура",
        "name": "PopulateEmployeeFieldsParams",
        "desc": "Структура параметров для заполнения полей сотрудника.",
        "fields": [
            {"code": "QComboBox* employmentRateCombo", "desc": "комбобокс коэффициента занятости"},
            {"code": "QComboBox* managerProject", "desc": "комбобокс проекта менеджера"},
            {"code": "QLineEdit* devLanguage", "desc": "поле языка разработчика"},
            {"code": "QLineEdit* devExperience", "desc": "поле опыта разработчика"},
            {"code": "QLineEdit* designerTool", "desc": "поле инструмента дизайнера"},
            {"code": "QLineEdit* designerProjects", "desc": "поле проектов дизайнера"},
            {"code": "QLineEdit* qaTestType", "desc": "поле типа тестирования QA"},
            {"code": "QLineEdit* qaBugs", "desc": "поле багов QA"},
            {"code": "QLabel* managerProjectLabel", "desc": "лейбл проекта менеджера"},
            {"code": "QLabel* devLanguageLabel", "desc": "лейбл языка программирования"},
            {"code": "QLabel* devExperienceLabel", "desc": "лейбл опыта разработчика"},
            {"code": "QLabel* designerToolLabel", "desc": "лейбл инструмента дизайнера"},
            {"code": "QLabel* designerProjectsLabel", "desc": "лейбл проектов дизайнера"},
            {"code": "QLabel* qaTestTypeLabel", "desc": "лейбл типа тестирования QA"},
            {"code": "QLabel* qaBugsLabel", "desc": "лейбл багов QA"},
            {"code": "std::shared_ptr<Employee> employee", "desc": "указатель на сотрудника"},
        ],
        "methods": [
            {"code": "(нет)", "desc": "структура только для передачи параметров"},
        ],
    })

    # 55. CreateEmployeeFromTypeParams
    sections.append({
        "kind": "Структура",
        "name": "CreateEmployeeFromTypeParams",
        "desc": "Структура параметров для создания сотрудника по типу.",
        "fields": [
            {"code": "const QString& employeeType", "desc": "тип сотрудника"},
            {"code": "int employeeId", "desc": "идентификатор сотрудника"},
            {"code": "const QString& name", "desc": "имя сотрудника"},
            {"code": "double salary", "desc": "зарплата сотрудника"},
            {"code": "const QString& department", "desc": "отдел сотрудника"},
            {"code": "QComboBox* employmentRateCombo", "desc": "комбобокс коэффициента занятости"},
            {"code": "QComboBox* managerProject", "desc": "комбобокс проекта менеджера"},
            {"code": "QLineEdit* devLanguage", "desc": "поле языка разработчика"},
            {"code": "QLineEdit* devExperience", "desc": "поле опыта разработчика"},
            {"code": "QLineEdit* designerTool", "desc": "поле инструмента дизайнера"},
            {"code": "QLineEdit* designerProjects", "desc": "поле проектов дизайнера"},
            {"code": "QLineEdit* qaTestType", "desc": "поле типа тестирования QA"},
            {"code": "QLineEdit* qaBugs", "desc": "поле багов QA"},
        ],
        "methods": [
            {"code": "(нет)", "desc": "структура только для передачи параметров"},
        ],
    })

    # 56. EditEmployeeParams
    sections.append({
        "kind": "Структура",
        "name": "EditEmployeeParams",
        "desc": "Структура параметров редактирования сотрудника.",
        "fields": [
            {"code": "QDialog* dialog", "desc": "диалог"},
            {"code": "Company* company", "desc": "компания"},
            {"code": "int employeeId", "desc": "идентификатор сотрудника"},
            {"code": "int& nextEmployeeId", "desc": "следующий ID сотрудника"},
            {"code": "const QLineEdit* nameEdit", "desc": "поле имени"},
            {"code": "const QLineEdit* salaryEdit", "desc": "поле зарплаты"},
            {"code": "const QLineEdit* deptEdit", "desc": "поле отдела"},
            {"code": "QComboBox* employmentRateCombo", "desc": "комбобокс коэффициента занятости"},
            {"code": "QComboBox* managerProject", "desc": "комбобокс проекта менеджера"},
            {"code": "QLineEdit* devLanguage", "desc": "поле языка разработчика"},
            {"code": "QLineEdit* devExperience", "desc": "поле опыта разработчика"},
            {"code": "QLineEdit* designerTool", "desc": "поле инструмента дизайнера"},
            {"code": "QLineEdit* designerProjects", "desc": "поле проектов дизайнера"},
            {"code": "QLineEdit* qaTestType", "desc": "поле типа тестирования QA"},
            {"code": "QLineEdit* qaBugs", "desc": "поле багов QA"},
            {"code": "const QString& currentType", "desc": "текущий тип сотрудника"},
        ],
        "methods": [
            {"code": "(нет)", "desc": "структура только для передачи параметров"},
        ],
    })

    # 57. EmployeeDialogHelper
    sections.append({
        "kind": "Класс",
        "name": "EmployeeDialogHelper",
        "desc": "Вспомогательный класс для работы с диалогами сотрудников.",
        "fields": [
            {"code": "(нет)", "desc": "содержит только статические методы"},
        ],
        "methods": [
            {"code": "static void createEmployeeDialog(const QDialog& dialog, QFormLayout* form, CreateEmployeeDialogFields& fields)", "desc": "создать диалог сотрудника"},
            {"code": "static void createEditEmployeeDialog(const QDialog& dialog, QFormLayout* form, std::shared_ptr<Employee> employee, CreateEditEmployeeDialogFields& fields)", "desc": "создать диалог редактирования сотрудника"},
            {"code": "static void showManagerFields(QLabel* managerProjectLabel, QComboBox* managerProject, bool show)", "desc": "показать/скрыть поля менеджера"},
            {"code": "static void showDeveloperFields(QLabel* devLanguageLabel, QLineEdit* devLanguage, QLabel* devExperienceLabel, QLineEdit* devExperience, bool show)", "desc": "показать/скрыть поля разработчика"},
            {"code": "static void showDesignerFields(QLabel* designerToolLabel, QLineEdit* designerTool, QLabel* designerProjectsLabel, QLineEdit* designerProjects, bool show)", "desc": "показать/скрыть поля дизайнера"},
            {"code": "static void showQaFields(QLabel* qaTestTypeLabel, QLineEdit* qaTestType, QLabel* qaBugsLabel, QLineEdit* qaBugs, bool show)", "desc": "показать/скрыть поля QA"},
            {"code": "static void setAllFieldsVisible(const SetAllFieldsVisibleParams& params)", "desc": "установить видимость всех полей"},
            {"code": "static void populateEmployeeFields(const PopulateEmployeeFieldsParams& params)", "desc": "заполнить поля данными сотрудника"},
            {"code": "static bool checkDuplicateEmployee(const QString& name, const Company* currentCompany)", "desc": "проверить дубликат сотрудника"},
            {"code": "static bool checkDuplicateEmployeeOnEdit(const QString& name, int excludeId, const Company* currentCompany)", "desc": "проверить дубликат при редактировании"},
            {"code": "static std::shared_ptr<Employee> createEmployeeFromType(const CreateEmployeeFromTypeParams& params)", "desc": "создать сотрудника по типу"},
        ],
    })

    # 58. AddEmployeeParams
    sections.append({
        "kind": "Структура",
        "name": "AddEmployeeParams",
        "desc": "Структура параметров добавления сотрудника.",
        "fields": [
            {"code": "QDialog* dialog", "desc": "диалог"},
            {"code": "Company* company", "desc": "компания"},
            {"code": "int& nextEmployeeId", "desc": "следующий ID сотрудника"},
            {"code": "const QLineEdit* nameEdit", "desc": "поле имени"},
            {"code": "const QLineEdit* salaryEdit", "desc": "поле зарплаты"},
            {"code": "const QLineEdit* deptEdit", "desc": "поле отдела"},
            {"code": "QComboBox* typeCombo", "desc": "комбобокс типа"},
            {"code": "QComboBox* employmentRateCombo", "desc": "комбобокс коэффициента занятости"},
            {"code": "QComboBox* managerProject", "desc": "комбобокс проекта менеджера"},
            {"code": "QLineEdit* devLanguage", "desc": "поле языка разработчика"},
            {"code": "QLineEdit* devExperience", "desc": "поле опыта разработчика"},
            {"code": "QLineEdit* designerTool", "desc": "поле инструмента дизайнера"},
            {"code": "QLineEdit* designerProjects", "desc": "поле проектов дизайнера"},
            {"code": "QLineEdit* qaTestType", "desc": "поле типа тестирования QA"},
            {"code": "QLineEdit* qaBugs", "desc": "поле багов QA"},
        ],
        "methods": [
            {"code": "(нет)", "desc": "структура только для передачи параметров"},
        ],
    })

    # 59. EmployeeDialogHandler
    sections.append({
        "kind": "Класс",
        "name": "EmployeeDialogHandler",
        "desc": "Обработчик диалогов сотрудников.",
        "fields": [
            {"code": "(нет)", "desc": "содержит только статические методы"},
        ],
        "methods": [
            {"code": "static bool processAddEmployee(const AddEmployeeParams& params)", "desc": "обработать добавление сотрудника"},
            {"code": "static bool processEditEmployee(const EditEmployeeParams& params)", "desc": "обработать редактирование сотрудника"},
        ],
    })

    # 60. TaskDialogHelper
    sections.append({
        "kind": "Класс",
        "name": "TaskDialogHelper",
        "desc": "Вспомогательный класс для работы с диалогами задач.",
        "fields": [
            {"code": "(нет)", "desc": "содержит только статические методы"},
        ],
        "methods": [
            {"code": "static void createAddTaskDialog(QDialog* dialog, QFormLayout* form, QLineEdit*& taskNameEdit, QComboBox*& taskTypeCombo, QLineEdit*& taskEstHoursEdit, QLineEdit*& priorityEdit)", "desc": "создать диалог добавления задачи"},
            {"code": "static bool validateAndAddTask(const QString& taskName, const QString& taskType, int taskEst, int priority, int projectId, const Company* company, QDialog* dialog)", "desc": "валидировать и добавить задачу"},
        ],
    })

    # 61. DialogHelper
    sections.append({
        "kind": "Класс",
        "name": "DialogHelper",
        "desc": "Вспомогательный класс для создания диалогов.",
        "fields": [
            {"code": "(нет)", "desc": "содержит только статические методы"},
        ],
        "methods": [
            {"code": "static void createHtmlDialog(QDialog* dialog, const QString& title, const QString& html, int minWidth = 800, int minHeight = 600)", "desc": "создать HTML-диалог"},
            {"code": "static void createTableDialog(QDialog* dialog, const QString& title, QTableWidget* table, int minWidth = 900, int minHeight = 600)", "desc": "создать диалог с таблицей"},
        ],
    })

    # 62. IdHelper
    sections.append({
        "kind": "Класс",
        "name": "IdHelper",
        "desc": "Вспомогательный класс для работы с идентификаторами.",
        "fields": [
            {"code": "(нет)", "desc": "содержит только статические методы"},
        ],
        "methods": [
            {"code": "static int findMaxEmployeeId(const std::vector<std::shared_ptr<Employee>>& employees)", "desc": "найти максимальный ID сотрудника"},
            {"code": "static int findMaxProjectId(const std::vector<Project>& projects)", "desc": "найти максимальный ID проекта"},
            {"code": "static int calculateNextId(int maxId)", "desc": "рассчитать следующий ID"},
        ],
    })

    # 63. FileHelper
    sections.append({
        "kind": "Класс",
        "name": "FileHelper",
        "desc": "Вспомогательный класс для работы с файлами.",
        "fields": [
            {"code": "(нет)", "desc": "содержит только статические методы"},
        ],
        "methods": [
            {"code": "static void clearAllDataFiles(QWidget* parent)", "desc": "очистить все файлы данных"},
        ],
    })

    # 64. HtmlGenerator
    sections.append({
        "kind": "Класс",
        "name": "HtmlGenerator",
        "desc": "Класс для генерации HTML-контента для отображения информации о проектах и сотрудниках.",
        "fields": [
            {"code": "(нет)", "desc": "содержит только статические методы"},
        ],
        "methods": [
            {"code": "static QString generateProjectDetailHtml(const Project& project, const Company* company)", "desc": "сгенерировать HTML деталей проекта"},
            {"code": "static QString generateProjectAssignmentsHtml(const Project& project, const Company* company)", "desc": "сгенерировать HTML назначений проекта"},
            {"code": "static QString generateEmployeeHistoryHtml(const Employee& employee, const Company* company, const std::vector<const Project*>& employeeProjects)", "desc": "сгенерировать HTML истории сотрудника"},
        ],
    })

    # 65. DisplayHelper
    sections.append({
        "kind": "Класс",
        "name": "DisplayHelper",
        "desc": "Вспомогательный класс для отображения данных в интерфейсе.",
        "fields": [
            {"code": "(нет)", "desc": "содержит только статические методы"},
        ],
        "methods": [
            {"code": "static void displayEmployees(QTableWidget* employeeTable, const Company* currentCompany, MainWindow* mainWindow)", "desc": "отобразить сотрудников в таблице"},
            {"code": "static void displayProjects(QTableWidget* projectTable, const Company* currentCompany, MainWindow* mainWindow)", "desc": "отобразить проекты в таблице"},
            {"code": "static void showCompanyInfo(QTextEdit* companyInfoText, const Company* currentCompany)", "desc": "показать информацию о компании"},
            {"code": "static void showStatistics(QTextEdit* statisticsText, const Company* currentCompany)", "desc": "показать статистику"},
            {"code": "static QString formatProjectInfo(const std::shared_ptr<const Employee>& employee, const Company* currentCompany)", "desc": "форматировать информацию о проекте"},
            {"code": "static QString formatTaskInfo(const std::shared_ptr<const Employee>& employee, const Company* currentCompany)", "desc": "форматировать информацию о задаче"},
            {"code": "static QString formatEmploymentRate(double rate)", "desc": "форматировать коэффициент занятости"},
        ],
    })

    # 66. ValidateEmployeeTypeFieldsParams
    sections.append({
        "kind": "Структура",
        "name": "ValidateEmployeeTypeFieldsParams",
        "desc": "Структура параметров валидации полей типа сотрудника.",
        "fields": [
            {"code": "const QString& employeeType", "desc": "тип сотрудника"},
            {"code": "QDialog* dialog", "desc": "диалог"},
            {"code": "QLineEdit* devLanguage", "desc": "поле языка разработчика"},
            {"code": "QLineEdit* devExperience", "desc": "поле опыта разработчика"},
            {"code": "QLineEdit* designerTool", "desc": "поле инструмента дизайнера"},
            {"code": "QLineEdit* designerProjects", "desc": "поле проектов дизайнера"},
            {"code": "QLineEdit* qaTestType", "desc": "поле типа тестирования QA"},
            {"code": "QLineEdit* qaBugs", "desc": "поле багов QA"},
        ],
        "methods": [
            {"code": "(нет)", "desc": "структура только для передачи параметров"},
        ],
    })

    # 67. EmployeeValidator
    sections.append({
        "kind": "Класс",
        "name": "EmployeeValidator",
        "desc": "Класс валидации данных сотрудников.",
        "fields": [
            {"code": "(нет)", "desc": "содержит только статические методы"},
        ],
        "methods": [
            {"code": "static bool validateEmployeeTypeFields(const ValidateEmployeeTypeFieldsParams& params)", "desc": "валидировать поля типа сотрудника"},
        ],
    })

    # 68. ValidationHelper
    sections.append({
        "kind": "Класс",
        "name": "ValidationHelper",
        "desc": "Вспомогательный класс для валидации входных данных.",
        "fields": [
            {"code": "(нет)", "desc": "содержит только статические методы"},
        ],
        "methods": [
            {"code": "static bool validateNonEmpty(const QString& value, const QString& fieldName, QDialog* dialog)", "desc": "проверить непустое значение"},
            {"code": "static bool validateDouble(const QString& text, double& result, double min, double max, const QString& fieldName, QDialog* dialog)", "desc": "валидировать дробное число"},
            {"code": "static bool validateInt(const QString& text, int& result, int min, int max, const QString& fieldName, QDialog* dialog)", "desc": "валидировать целое число"},
            {"code": "static bool validateDateRange(const QDate& startDate, const QDate& endDate, QDialog* dialog)", "desc": "валидировать диапазон дат"},
        ],
    })

    # 69. ProjectHelper
    sections.append({
        "kind": "Класс",
        "name": "ProjectHelper",
        "desc": "Вспомогательный класс для работы с проектами.",
        "fields": [
            {"code": "(нет)", "desc": "содержит только статические методы"},
        ],
        "methods": [
            {"code": "static void populateProjectTasksTable(QTableWidget* table, const Project& project, MainWindow* mainWindow)", "desc": "заполнить таблицу задач проекта"},
            {"code": "static void clearProjectAllocatedHoursIfNoEmployees(const Company* company, int projectId)", "desc": "очистить выделенные часы проекта если нет сотрудников"},
            {"code": "static bool hasAssignedEmployees(const Company* company, int projectId)", "desc": "проверить наличие назначенных сотрудников"},
        ],
    })

    # 70. TaskAssignmentHelper
    sections.append({
        "kind": "Класс",
        "name": "TaskAssignmentHelper",
        "desc": "Вспомогательный класс для работы с назначениями задач.",
        "fields": [
            {"code": "(нет)", "desc": "содержит только статические методы"},
        ],
        "methods": [
            {"code": "static QString getExpectedRoleForProjectPhase(const QString& projectPhase)", "desc": "получить ожидаемую роль для фазы проекта"},
            {"code": "static void populateEmployeeCombo(QComboBox* employeeCombo, const Company* company, int projectId, const QString& projectPhase, int& matchingCount)", "desc": "заполнить комбобокс сотрудников"},
            {"code": "static void setupTaskCombo(QComboBox* taskCombo, const std::vector<Task>& tasks, int pendingTaskId)", "desc": "настроить комбобокс задач"},
            {"code": "static void setupHoursEdit(QLineEdit* hoursEdit, const QComboBox* taskCombo, const QComboBox* employeeCombo, const std::vector<Task>& tasks, const Company* company)", "desc": "настроить поле часов"},
            {"code": "static void setupEmployeeComboUpdate(QComboBox* employeeCombo, const QComboBox* taskCombo, QLabel* taskInfoLabel, const Company* company, int projectId, const QString& projectPhase)", "desc": "настроить обновление комбобокса сотрудников"},
        ],
    })

    # 71. ActionButtonHelper
    sections.append({
        "kind": "Класс",
        "name": "ActionButtonHelper",
        "desc": "Вспомогательный класс для создания кнопок действий в таблицах.",
        "fields": [
            {"code": "(нет)", "desc": "содержит только статические методы"},
        ],
        "methods": [
            {"code": "static QWidget* createEmployeeActionButtons(QTableWidget* table, int rowIndex, MainWindow* mainWindow)", "desc": "создать кнопки действий для сотрудника"},
            {"code": "static QWidget* createProjectActionButtons(QTableWidget* table, int rowIndex, MainWindow* mainWindow, bool includeAddTask = true)", "desc": "создать кнопки действий для проекта"},
        ],
    })

    # 72. UIStyleHelper
    sections.append({
        "kind": "Класс",
        "name": "UIStyleHelper",
        "desc": "Вспомогательный класс для стилей интерфейса.",
        "fields": [
            {"code": "(нет)", "desc": "содержит только статические методы"},
        ],
        "methods": [
            {"code": "static QString getMainWindowStylesheet()", "desc": "получить таблицу стилей главного окна"},
            {"code": "static QString getCompanyComboBoxStylesheet()", "desc": "получить таблицу стилей комбобокса компании"},
        ],
    })

    return sections


def build_docx(output_path: str):
    doc = Document()

    # Заголовок раздела
    make_header_paragraph(doc, "3.1 Описание программных модулей")

    sections = build_sections()
    for idx, sec in enumerate(sections, start=1):
        make_class_heading(doc, idx, sec["kind"], sec["name"])
        make_normal_paragraph(doc, sec["desc"])
        add_bullet_block(doc, "Поля:", sec["fields"])
        add_bullet_block(doc, "Методы:", sec["methods"])

    doc.save(output_path)


if __name__ == "__main__":
    build_docx("Раздел_3_1_классы_и_структуры.docx")
    print("Готово: Раздел_3_1_классы_и_структуры.docx")

